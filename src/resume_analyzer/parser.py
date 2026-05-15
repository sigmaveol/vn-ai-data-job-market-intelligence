"""Lightweight deterministic resume parser for PDF, DOCX, and TXT files."""
from __future__ import annotations

import re
import tempfile
from pathlib import Path
from typing import BinaryIO


SECTION_ALIASES = {
    "skills": [
        "skills", "technical skills", "kỹ năng", "ky nang", "core competencies",
        "technologies", "tools",
    ],
    "experience": [
        "experience", "work experience", "employment", "kinh nghiệm", "kinh nghiem",
        "professional experience",
    ],
    "education": [
        "education", "học vấn", "hoc van", "academic background", "qualification",
    ],
    "projects": [
        "projects", "project experience", "dự án", "du an", "portfolio",
    ],
    "certifications": [
        "certifications", "certificates", "chứng chỉ", "chung chi",
    ],
    "summary": [
        "summary", "profile", "objective", "career objective", "giới thiệu",
    ],
}


class ResumeParser:
    """Extract plain text and coarse resume sections without hallucination."""

    def parse(self, file_path: Path) -> str:
        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            return self._parse_pdf(file_path)
        if suffix in (".doc", ".docx"):
            return self._parse_docx(file_path)
        if suffix in (".txt", ".text"):
            return self._parse_txt(file_path)
        raise ValueError(f"Unsupported format: {suffix}")

    def parse_uploaded_file(self, uploaded_file) -> str:
        suffix = Path(uploaded_file.name).suffix.lower()
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(uploaded_file.getbuffer())
            tmp_path = Path(tmp.name)
        try:
            return self.parse(tmp_path)
        finally:
            try:
                tmp_path.unlink(missing_ok=True)
            except Exception:
                pass

    def parse_bytes(self, file_obj: BinaryIO, suffix: str) -> str:
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(file_obj.read())
            tmp_path = Path(tmp.name)
        try:
            return self.parse(tmp_path)
        finally:
            tmp_path.unlink(missing_ok=True)

    def _parse_pdf(self, path: Path) -> str:
        try:
            import pdfplumber
        except ImportError as exc:
            raise RuntimeError("pdfplumber is required for PDF resume parsing.") from exc

        chunks = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                chunks.append(page.extract_text() or "")
        return self.clean_text("\n".join(chunks))

    def _parse_docx(self, path: Path) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("python-docx is required for DOCX resume parsing.") from exc

        doc = Document(path)
        chunks = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                chunks.append(" | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()))
        return self.clean_text("\n".join(chunks))

    def _parse_txt(self, path: Path) -> str:
        for encoding in ("utf-8-sig", "utf-8", "cp1258", "latin-1"):
            try:
                return self.clean_text(path.read_text(encoding=encoding))
            except UnicodeDecodeError:
                continue
        return self.clean_text(path.read_text(errors="ignore"))

    @staticmethod
    def clean_text(text: str) -> str:
        text = str(text or "").replace("\x00", " ")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def extract_sections(self, text: str) -> dict:
        lines = [line.strip() for line in str(text or "").splitlines() if line.strip()]
        section_for_line = {}
        for idx, line in enumerate(lines):
            normalized = self._normalize_header(line)
            for section, aliases in SECTION_ALIASES.items():
                if normalized in {self._normalize_header(alias) for alias in aliases}:
                    section_for_line[idx] = section
                    break

        sections = {key: "" for key in SECTION_ALIASES}
        if not lines:
            return sections

        markers = sorted(section_for_line.items())
        if not markers:
            sections["summary"] = "\n".join(lines[:12])
            return sections

        for pos, (line_idx, section) in enumerate(markers):
            next_idx = markers[pos + 1][0] if pos + 1 < len(markers) else len(lines)
            sections[section] = "\n".join(lines[line_idx + 1:next_idx]).strip()

        if not sections["summary"]:
            first_marker = markers[0][0]
            sections["summary"] = "\n".join(lines[:first_marker]).strip()
        return sections

    @staticmethod
    def _normalize_header(text: str) -> str:
        text = str(text or "").lower().strip(":：-–— ")
        text = re.sub(r"[^a-zÀ-ỹ ]+", " ", text)
        text = re.sub(r"\s+", " ", text).strip()
        return text

    def extract_experience_years(self, text: str) -> float | None:
        text = str(text or "").lower()
        patterns = [
            r"(\d+(?:\.\d+)?)\+?\s*(?:years?|yrs?)",
            r"(\d+(?:\.\d+)?)\+?\s*(?:năm|nam)\s*(?:kinh nghiệm|kinh nghiem)?",
        ]
        values = []
        for pattern in patterns:
            values.extend(float(match) for match in re.findall(pattern, text))
        return max(values) if values else None
